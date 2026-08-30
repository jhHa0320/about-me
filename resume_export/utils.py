import os
import re
import tempfile
from contextlib import contextmanager
from datetime import date
from io import BytesIO
from pathlib import Path

import docx
from docx.shared import Inches, Pt
from django.conf import settings
from django.template.loader import render_to_string
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from xhtml2pdf import pisa

from portfolio.models import Activity, Career, Leadership, Profile, Project
from portfolio.templatetags.portfolio_extras import bullets

from .models import ResumeExportConfig

# 한글 PDF 텍스트용 폰트를 등록한다. reportlab 내장 CID 폰트(HYGothic-Medium)는
# 한글은 문제없지만 영문/숫자 글자폭이 어색하게 벌어져서, 한글·영문이 같이
# 자연스러운 나눔고딕 TTF 를 직접 번들링해서 쓴다 (resume_export/static/.../fonts/).
_FONT_DIR = Path(__file__).resolve().parent / "static" / "resume_export" / "fonts"
_KOREAN_FONT = "NanumGothic"
_KOREAN_FONT_BOLD = "NanumGothic-Bold"
pdfmetrics.registerFont(TTFont(_KOREAN_FONT, str(_FONT_DIR / "NanumGothic-Regular.ttf")))
pdfmetrics.registerFont(TTFont(_KOREAN_FONT_BOLD, str(_FONT_DIR / "NanumGothic-Bold.ttf")))
pdfmetrics.registerFontFamily(
    _KOREAN_FONT, normal=_KOREAN_FONT, bold=_KOREAN_FONT_BOLD,
    italic=_KOREAN_FONT, boldItalic=_KOREAN_FONT_BOLD,
)

SKILL_DOMAIN_ORDER = [
    ("LANGUAGE", "Language"),
    ("DATA_SCIENCE", "Data Science"),
    ("AI", "AI"),
    ("SECURITY", "Security"),
    ("BACKEND", "Backend"),
    ("ETC", "기타"),
]

# 파일명에 쓰면 안 되는 문자들을 걷어낸다 (Windows/Linux 공통 금지 문자 기준).
_UNSAFE_FILENAME_CHARS = re.compile(r'[\\/:*?"<>|]+')


def get_export_config():
    """제외 목록 싱글톤을 가져온다 (admin 에서 아직 안 만졌으면 자동으로 하나 만든다)."""
    config, _ = ResumeExportConfig.objects.get_or_create(pk=1)
    return config


def _media_path(url_path):
    relative = url_path[len(settings.MEDIA_URL):].lstrip("/")
    candidate = os.path.join(str(settings.MEDIA_ROOT), relative)
    return candidate if os.path.isfile(candidate) else None


def _resolve_media_or_static_path(uri):
    """/static/ 또는 /media/ 로 시작하는 URL 을 실제 디스크 경로로 바꾼다.

    PDF/DOCX 생성 모두 HTTP 요청 없이 로컬에서 바로 이뤄지므로, 이미지·폰트를
    넣으려면 URL 이 아니라 실제 파일 경로가 필요하다. static 쪽은 Django 의
    staticfiles finder 를 그대로 써서 collectstatic 을 아직 안 돌린 로컬
    개발 환경에서도, 운영(collectstatic 이후)에서도 똑같이 동작하게 한다.
    """
    if uri.startswith(settings.MEDIA_URL):
        return _media_path(uri)
    if uri.startswith(settings.STATIC_URL):
        from django.contrib.staticfiles import finders

        relative = uri[len(settings.STATIC_URL):].lstrip("/")
        return finders.find(relative)
    return None


def resume_pdf_link_callback(uri, rel):
    """xhtml2pdf 가 <img src="..."> 등을 실제 파일 경로로 바꿀 때 쓰는 콜백."""
    if not (uri.startswith(settings.MEDIA_URL) or uri.startswith(settings.STATIC_URL)):
        return uri  # http(s):// 등 절대 URI 는 그대로 xhtml2pdf 에 맡긴다.
    return _resolve_media_or_static_path(uri) or uri


def _skill_groups():
    from portfolio.models import Skill

    skills = Skill.objects.all().order_by("name")
    by_domain = {}
    for skill in skills:
        by_domain.setdefault(skill.domain, []).append(skill)

    return [
        {"key": key, "label": label, "skills": by_domain.get(key, [])}
        for key, label in SKILL_DOMAIN_ORDER
        if by_domain.get(key)
    ]


def build_resume_context() -> dict:
    """내보내기 설정의 제외 목록을 제외한 전체 이력서/포트폴리오 데이터 템플릿 컨텍스트를 구성합니다.

    Returns:
        dict: PDF 및 DOCX 렌더링에 사용될 모델 객체 및 메타 데이터 딕셔너리.

    Rationale:
        관리자가 제외 목록(ResumeExportConfig)을 빈 값으로 둘 경우 사이트의 최신 전체 데이터가 자동으로 내보내기에 반영되며,
        특정 데이터(예: 비공개 프로젝트)만 제외하도록 세밀히 제어할 수 있는 동적 템플릿 구조를 구축했습니다.
    """
    config = get_export_config()
    profile = Profile.objects.first()

    return {
        "profile": profile,
        "educations": profile.educations.all() if profile else [],
        "skill_groups": _skill_groups(),
        "projects": (
            Project.objects.filter(is_active=True)
            .exclude(pk__in=config.excluded_projects.all())
            .prefetch_related("tech_stacks")
        ),
        "careers": Career.objects.exclude(pk__in=config.excluded_careers.all()),
        "leaderships": Leadership.objects.exclude(pk__in=config.excluded_leaderships.all()),
        "certifications": Activity.objects.filter(type="CERTIFICATION")
            .exclude(pk__in=config.excluded_activities.all()),
        "awards": Activity.objects.filter(type="AWARD")
            .exclude(pk__in=config.excluded_activities.all()),
        "activities": Activity.objects.filter(type="ACTIVITY")
            .exclude(pk__in=config.excluded_activities.all()),
        "generated_on": date.today(),
    }


@contextmanager
def _windows_tempfile_reopen_fix():
    """Windows 환경에서 xhtml2pdf의 TTF 폰트 임시 파일 권한 에러를 우회하는 컨텍스트 매니저.

    Rationale:
        NOTE: xhtml2pdf가 @font-face로 임베드한 TTF 폰트를 임시 파일로 복사한 뒤 reportlab이 해당 파일을 재열 때,
        Windows OS에서는 열려 있는 NamedTemporaryFile의 재오픈을 차단하여 PermissionError가 발생하는 이슈가 존재합니다.
        Linux(PythonAnywhere)에는 해당 문제가 없으므로 os.name == 'nt' 일 때만 파일 delete=False로 패치하여 우회하도록 설계했습니다.
    """
    if os.name != "nt":
        yield
        return

    original = tempfile.NamedTemporaryFile

    def patched(*args, **kwargs):
        kwargs["delete"] = False
        return original(*args, **kwargs)

    tempfile.NamedTemporaryFile = patched
    try:
        yield
    finally:
        tempfile.NamedTemporaryFile = original


def render_resume_pdf_bytes() -> bytes:
    """포트폴리오 HTML 템플릿을 xhtml2pdf 엔진을 이용해 PDF 바이너리 바이트로 렌더링합니다.

    Returns:
        bytes: 생성된 PDF 파일 바이너리 데이터.

    Raises:
        RuntimeError: xhtml2pdf 변환 과정 중 에러 발생 시 예외 throw.

    Rationale:
        ReportLab의 낮은 수준 API로 PDF 레이아웃을 직접 하드코딩하는 대신 HTML/CSS 템플릿 기반으로 xhtml2pdf를 사용함으로써
        스타일 수정이 용이하고 유지보수성이 뛰어난 PDF 내보내기 구조를 달성했습니다.
    """
    html = render_to_string("resume_export/resume_pdf.html", build_resume_context())
    buffer = BytesIO()
    with _windows_tempfile_reopen_fix():
        result = pisa.CreatePDF(
            src=html, dest=buffer, link_callback=resume_pdf_link_callback, encoding="utf-8"
        )
    if result.err:
        raise RuntimeError(f"PDF 생성 실패 (xhtml2pdf err={result.err})")
    return buffer.getvalue()


def build_resume_filename(profile, ext: str, suffix: str = "전체") -> str:
    """안전한 이력서 다운로드 파일명을 생성합니다.

    Args:
        profile (Profile): 프로필 모델 객체.
        ext (str): 확장자 ('pdf' 또는 'docx').
        suffix (str, optional): 파일명 접미사. Defaults to "전체".

    Returns:
        str: 정형화된 다운로드 파일명 문자열 (예: '홍길동_포트폴리오_전체_20260830.pdf').
    """
    safe_name = _UNSAFE_FILENAME_CHARS.sub("", profile.name if profile else "").strip() or "포트폴리오"
    today = date.today().strftime("%Y%m%d")
    return f"{safe_name}_포트폴리오_{suffix}_{today}.{ext}"


def _add_heading(document, text):
    heading = document.add_heading(text, level=1)
    heading.runs[0].font.size = Pt(14)


def render_resume_docx_bytes() -> bytes:
    """python-docx 라이브러리를 이용하여 포트폴리오 데이터를 워드 DOCX 바이너리로 생성합니다.

    Returns:
        bytes: 생성된 DOCX 문서 바이너리 바이트.

    Rationale:
        Admin 사용자가 자유롭게 이력서 내용을 편집할 수 있도록 docx 객체를 구성하여 제공합니다.
    """
    context = build_resume_context()
    profile = context["profile"]
    document = docx.Document()

    title = document.add_heading(profile.name if profile else "", level=0)
    title.runs[0].font.size = Pt(22)

    if profile and profile.headline:
        document.add_paragraph(profile.headline)

    if profile and profile.avatar_url:
        image_path = _resolve_media_or_static_path(profile.avatar_url)
        if image_path:
            document.add_picture(image_path, width=Inches(1))

    contact_bits = []
    if profile and profile.show_email_address:
        contact_bits.append(f"Email: {profile.email}")
    if profile and profile.github_url:
        contact_bits.append(f"GitHub: {profile.github_url}")
    if profile and profile.show_birthdate:
        contact_bits.append(f"생년월일: {profile.birthdate:%Y.%m.%d}")
    if profile and profile.english_score:
        contact_bits.append(profile.english_score)
    if contact_bits:
        document.add_paragraph(" | ".join(contact_bits))

    if profile and profile.introduction:
        document.add_paragraph(profile.introduction)

    if context["skill_groups"]:
        _add_heading(document, "기술 스택")
        for group in context["skill_groups"]:
            names = ", ".join(skill.name for skill in group["skills"])
            p = document.add_paragraph()
            p.add_run(f"{group['label']}: ").bold = True
            p.add_run(names)

    if context["projects"]:
        _add_heading(document, "프로젝트")
        for project in context["projects"]:
            p = document.add_paragraph()
            title_text = project.title
            if project.key_result:
                title_text += f" [{project.key_result}]"
            p.add_run(f"{title_text}  ").bold = True
            p.add_run(project.period)

            meta = f"역할: {project.role}"
            tech_names = ", ".join(t.name for t in project.tech_stacks.all())
            if tech_names:
                meta += f" | 기술: {tech_names}"
            document.add_paragraph(meta)

            document.add_paragraph(project.description)
            for point in bullets(project.outcome):
                document.add_paragraph(point, style="List Bullet")

    if context["careers"] or context["educations"]:
        _add_heading(document, "경력 / 학력")
        for career in context["careers"]:
            p = document.add_paragraph()
            p.add_run(f"{career.organization} - {career.role}").bold = True
            document.add_paragraph(career.period)
            document.add_paragraph(career.description)
        for edu in context["educations"]:
            p = document.add_paragraph()
            p.add_run(f"{edu.school} ({edu.status})").bold = True
            document.add_paragraph(edu.period)

    if context["leaderships"]:
        _add_heading(document, "리더십 및 활동")
        for lead in context["leaderships"]:
            p = document.add_paragraph()
            p.add_run(f"{lead.title} - {lead.organization} ({lead.role})").bold = True
            document.add_paragraph(lead.period)
            document.add_paragraph(lead.description)

    if context["certifications"] or context["awards"] or context["activities"]:
        _add_heading(document, "자격증 / 수상 / 대외활동")
        for item in context["certifications"]:
            document.add_paragraph(f"[자격증] {item.title} - {item.organization}")
        for item in context["awards"]:
            document.add_paragraph(f"[수상] {item.title} - {item.organization}")
        for item in context["activities"]:
            document.add_paragraph(f"[대외활동] {item.title} - {item.organization}")

    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()
